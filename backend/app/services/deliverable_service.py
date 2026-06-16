"""Deal-scoped deliverable generation on top of SQLite + local storage.

Pulls analyses + document text via a SQLAlchemy session, asks the task-routed
LLM to produce a Markdown document, renders it to .docx via python-docx, and
saves the result to the ``deliverables`` bucket on local disk. Returns a
short-lived HMAC-signed download URL.
"""

from __future__ import annotations

import io
import json
import uuid
from datetime import UTC, datetime
from typing import Any
from uuid import UUID

import structlog
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import Settings
from app.db.models import Analysis, Document, Meeting
from app.llm.router import TASK_IC_MEMO, LLMRouter

logger = structlog.get_logger(__name__)

DELIVERABLES_BUCKET = "deliverables"

_TYPE_PROMPTS: dict[str, str] = {
    "investment_memo": (
        "You are preparing an investment memo for a private-equity / investment-"
        "banking team. Using the deal context below, produce a memo with these "
        "Markdown sections in order: Executive Summary, Company Overview, "
        "Financial Analysis, Market & Competition, Key Risks, Valuation, "
        "Recommendation. Be specific and cite evidence from the transcripts / "
        "documents where possible. Return Markdown only — no preamble."
    ),
    "ic_presentation": (
        "You are outlining an IC presentation deck. Using the deal context, "
        "produce a Markdown document with one top-level section per slide: "
        "Deal Overview, Investment Thesis (1), Investment Thesis (2), Market & "
        "Competitive Landscape, Financial Summary, Projections, Key Risks & "
        "Mitigants, Terms & Next Steps. Each section gets 3–6 bullet points. "
        "Markdown only — no preamble."
    ),
    "financial_model": (
        "You are drafting the assumptions section of a financial model. Using "
        "the deal context, produce a Markdown document titled 'Model "
        "Assumptions' with sections: Revenue Drivers, Margin Assumptions, "
        "Operating Costs, Capital Structure, Scenario Ranges (bear / base / "
        "bull). Return Markdown only — no preamble."
    ),
}

_TITLES: dict[str, str] = {
    "investment_memo": "Investment Memo",
    "ic_presentation": "IC Presentation",
    "financial_model": "Financial Model",
}


class DeliverableService:
    def __init__(
        self,
        session: Session,
        settings: Settings,
        llm_router: LLMRouter,
    ) -> None:
        self.session = session
        self.settings = settings
        self.llm_router = llm_router

    async def generate(
        self,
        deal_id: UUID,
        deliverable_type: str,
    ) -> dict[str, Any]:
        context = await self._gather_context(deal_id)
        prompt = _TYPE_PROMPTS.get(deliverable_type, _TYPE_PROMPTS["investment_memo"])
        markdown = await self._render_markdown(prompt, context)
        docx_bytes = _markdown_to_docx(
            markdown, title=_TITLES.get(deliverable_type, "Deliverable")
        )

        from app.storage.local import make_signed_url, save_bytes

        file_key = f"{deal_id}/{uuid.uuid4()}.docx"
        save_bytes(DELIVERABLES_BUCKET, file_key, docx_bytes)
        download_url = self.settings.public_api_url + make_signed_url(
            DELIVERABLES_BUCKET, file_key
        )

        logger.info(
            "deliverable_generated",
            deal_id=str(deal_id),
            type=deliverable_type,
            file_key=file_key,
            markdown_chars=len(markdown),
        )

        return {
            "id": str(uuid.uuid4()),
            "deal_id": str(deal_id),
            "title": (
                f"{_TITLES.get(deliverable_type, 'Deliverable')} - "
                f"{datetime.now(UTC).strftime('%Y-%m-%d')}"
            ),
            "deliverable_type": deliverable_type,
            "file_format": "docx",
            "file_key": file_key,
            "status": "ready",
            "download_url": download_url,
            "created_at": datetime.now(UTC).isoformat(),
        }

    async def _gather_context(self, deal_id: UUID) -> str:
        parts: list[str] = [f"Deal ID: {deal_id}"]

        meetings = self.session.execute(
            select(Meeting).where(Meeting.deal_id == str(deal_id))
        ).scalars().all()
        parts.append(f"Meetings on record: {len(meetings)}")

        if meetings:
            meeting_ids = [m.id for m in meetings]
            analyses = self.session.execute(
                select(Analysis)
                .where(Analysis.meeting_id.in_(meeting_ids))
                .where(Analysis.status == "completed")
            ).scalars().all()
            for a in analyses:
                parts.append(
                    f"\n--- Analysis ({a.call_type or '?'} v{a.version or 1}) ---"
                )
                result = a.structured_output or {}
                parts.append(json.dumps(result, indent=2)[:4000])

        docs = self.session.execute(
            select(Document).where(Document.deal_id == str(deal_id))
        ).scalars().all()
        for doc in docs:
            text = doc.extracted_text or ""
            if not text:
                continue
            parts.append(
                f"\n--- Document: {doc.title or '?'} ({doc.document_type or '?'}) ---"
            )
            parts.append(text[:4000])

        return "\n".join(parts)[:40000]

    async def _render_markdown(self, system_prompt: str, context: str) -> str:
        response = await self.llm_router.complete(
            task_type=TASK_IC_MEMO,
            system_prompt=system_prompt,
            user_prompt=f"Deal context:\n\n{context}",
            max_tokens=4096,
            temperature=0.4,
        )
        return response.content.strip() or "# Deliverable\n\n_Empty LLM response._"


def _markdown_to_docx(markdown: str, title: str) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(title, level=0)

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
